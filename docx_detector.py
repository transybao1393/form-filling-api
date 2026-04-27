"""
Docx form detector.

Walks the body of a .docx in document order, finds empty table cells, and
classifies each one as a fillable field. Two layouts are recognised:

  1. Left-label rows: cell to the left of the empty cell holds the label
     (e.g. respondent-info tables: "Organization Name | <empty>").

  2. Answer rows: cell to the left says "Answer" (or a generic header
     variant), and the *previous row* carries the question — typically
     [QID, question_text] for ESG-style section tables.

The output JSON mirrors the shape of fields_enriched.json from the PDF
pipeline, so flatlist_adapter.match_fields_to_items() can consume it
without modification. Each field carries a `docx_locator` for the writer.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

from field_normalizer import BUILTIN_SECTION_PATTERNS
from flatlist_adapter import _is_generic_left_label, _QNUM_RE


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _iter_body(doc):
    """Yield ('p', paragraph) or ('tbl', table_index) in document order."""
    body = doc.element.body
    table_idx = 0
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield "p", child
        elif tag == "tbl":
            yield "tbl", table_idx
            table_idx += 1


def _para_text(p_element) -> str:
    parts = []
    for t in p_element.iter(f"{{{W_NS}}}t"):
        if t.text:
            parts.append(t.text)
    return "".join(parts).strip()


def _cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _slugify(text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")
    return slug or "field"


# "Section A — Governance & General Disclosures", "Section B – Economic …"
_SECTION_HEADER_RE = re.compile(
    r"^section\s+[A-Z0-9]+\s*[—–\-:]\s*(.+)$", re.I,
)


def _heading_to_section(text: str) -> str | None:
    """Map a heading-like paragraph to a canonical section slug."""
    if not text:
        return None
    compact = re.sub(r"\s+", "", text).lower()
    for pat, name in BUILTIN_SECTION_PATTERNS:
        target = re.sub(r"\\s\+", "", pat.pattern).lower()
        target = re.sub(r"[\^\$\\]", "", target)
        if target and target in compact:
            return name
    # "Section B — Economic Performance" → derive from descriptor.
    m = _SECTION_HEADER_RE.match(text.strip())
    if m:
        descriptor = m.group(1).strip()
        # Re-check built-ins on the descriptor alone
        compact_d = re.sub(r"\s+", "", descriptor).lower()
        for pat, name in BUILTIN_SECTION_PATTERNS:
            target = re.sub(r"\\s\+", "", pat.pattern).lower()
            target = re.sub(r"[\^\$\\]", "", target)
            if target and target in compact_d:
                return name
        # Fallback: slugify the descriptor (strip "Performance"/"Disclosures").
        descriptor = re.sub(
            r"\b(performance|disclosures?|information)\b",
            "", descriptor, flags=re.I,
        )
        return _slugify(descriptor) or None
    return None


def _looks_like_heading(text: str) -> bool:
    """A short line matching a known section keyword or 'Section X — …' form."""
    if not text or len(text) > 100:
        return False
    return _heading_to_section(text) is not None


def detect_docx_fields_to_json(docx_path: str | Path,
                               out_path: str | Path) -> dict:
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))
    tables = list(doc.tables)

    # Pass 1: walk the body to map each table index → its preceding section.
    section_for_table: dict[int, str] = {}
    current_section = "general"
    for kind, payload in _iter_body(doc):
        if kind == "p":
            text = _para_text(payload)
            sec = _heading_to_section(text) if _looks_like_heading(text) else None
            if sec:
                current_section = sec
        elif kind == "tbl":
            section_for_table[payload] = current_section

    fields: list[dict] = []
    counts: dict[str, int] = {}

    # Pass 2: walk each table and find empty cells.
    for ti, table in enumerate(tables):
        section = section_for_table.get(ti, "general")
        rows = table.rows
        for ri, row in enumerate(rows):
            cells = row.cells
            for ci, cell in enumerate(cells):
                if _cell_text(cell):
                    continue
                # Need a cell to the left to anchor a label.
                if ci == 0:
                    continue
                left_cell_text = _cell_text(cells[ci - 1])
                if not left_cell_text:
                    continue

                question_text = ""
                question_number: str | None = None

                # If left cell is a generic "Answer" header, look at the
                # row above for [QID | question_text].
                if _is_generic_left_label(left_cell_text) and ri > 0:
                    above = rows[ri - 1].cells
                    if len(above) >= 2:
                        qn_cell = _cell_text(above[0])
                        qt_cell = _cell_text(above[1])
                        m = _QNUM_RE.match(qn_cell)
                        if m:
                            question_number = m.group(1).upper()
                        # Allow either column to carry the prose.
                        question_text = (qt_cell or qn_cell).strip()

                # Pick a label: real left-label wins; otherwise question text.
                if question_text and _is_generic_left_label(left_cell_text):
                    label = question_text[:80]
                else:
                    label = left_cell_text

                # Build canonical key.
                base = _slugify(label) or f"field_{ti}_{ri}_{ci}"
                if section and section != "general":
                    base = f"{section}__{base}"
                counts[base] = counts.get(base, 0) + 1
                canonical_key = base if counts[base] == 1 else f"{base}_{counts[base]}"

                field = {
                    "field_id": f"table{ti}_row{ri}_col{ci}",
                    "label": label,
                    "page": 1,
                    "bbox": [ti, ri, ci, 0],
                    "page_width": 0,
                    "page_height": 0,
                    "confidence": 1.0,
                    "strategy": "docx_table",
                    "field_type": "text",
                    "acroform_name": None,
                    "section": section,
                    "canonical_key": canonical_key,
                    "left_label": left_cell_text,
                    "docx_locator": {"table": ti, "row": ri, "col": ci},
                }
                if question_text:
                    field["question_text"] = question_text
                if question_number:
                    field["question_number"] = question_number
                fields.append(field)

    out = {
        "source_docx": str(docx_path),
        "num_fields": len(fields),
        "fields_by_strategy": {"docx_table": len(fields)},
        "fields": fields,
    }
    Path(out_path).write_text(json.dumps(out, indent=2, ensure_ascii=False))
    return out


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("-o", "--output", default="docx_fields.json")
    args = ap.parse_args()
    data = detect_docx_fields_to_json(args.docx, args.output)
    print(f"Detected {data['num_fields']} fields → {args.output}")
    for f in data["fields"][:10]:
        print(f"  [{f['section']:20s}] {f['canonical_key']:40s}"
              f"  ← '{f['label'][:50]}'")
