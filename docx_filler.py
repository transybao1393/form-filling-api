"""
Docx form filler.

Given:
  - The source .docx (read-only),
  - A fields JSON whose entries each carry a `docx_locator`
    {table, row, col} pointing at the empty cell to fill,
  - A flat data map { canonical_key: answer_text },

writes the answer text into each target cell and saves a new .docx.

Mirrors form_filler.fill_form()'s return shape so run_pipeline can
report the same way for both formats.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document


def _set_cell_text(cell, text: str) -> None:
    """
    Replace cell content with `text`, preserving the formatting of the
    cell's first existing run when possible (so filled answers inherit
    the document's body font/size).
    """
    text = "" if text is None else str(text)

    # python-docx exposes cell.paragraphs; we keep the first paragraph,
    # clear it, drop any extras, then write into the first paragraph.
    paragraphs = cell.paragraphs
    first_para = paragraphs[0]

    # Remove any extra paragraphs after the first
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    # Capture style of the first existing run, if any
    template_run = first_para.runs[0] if first_para.runs else None

    # Clear runs in the first paragraph
    for run in list(first_para.runs):
        run._element.getparent().remove(run._element)

    new_run = first_para.add_run(text)
    if template_run is not None:
        src_rpr = template_run._element.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
        )
        if src_rpr is not None:
            new_run._element.insert(0, src_rpr)


def fill_docx(
    source_docx: str | Path,
    fields_json: str | Path,
    user_data: dict,
    output_docx: str | Path,
) -> dict:
    fields_data = json.loads(Path(fields_json).read_text())
    fields = fields_data["fields"]

    doc = Document(str(source_docx))
    tables = list(doc.tables)

    num_filled = 0
    num_missing = 0
    missing_keys: list[str] = []

    for f in fields:
        key = f["canonical_key"]
        if key not in user_data:
            num_missing += 1
            missing_keys.append(key)
            continue
        answer = user_data[key]
        if answer is None or answer == "":
            num_missing += 1
            missing_keys.append(key)
            continue

        loc = f.get("docx_locator")
        if not loc:
            num_missing += 1
            continue
        ti, ri, ci = loc["table"], loc["row"], loc["col"]
        if ti >= len(tables):
            num_missing += 1
            continue
        table = tables[ti]
        if ri >= len(table.rows):
            num_missing += 1
            continue
        row = table.rows[ri]
        if ci >= len(row.cells):
            num_missing += 1
            continue
        cell = row.cells[ci]

        _set_cell_text(cell, str(answer))
        num_filled += 1

    Path(output_docx).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))

    return {
        "num_filled": num_filled,
        "num_missing": num_missing,
        "missing_keys": missing_keys[:50],
        "output_pdf": str(output_docx),  # name kept for run_pipeline parity
    }


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True)
    ap.add_argument("--fields", required=True)
    ap.add_argument("--data", required=True)
    ap.add_argument("-o", "--output", default="filled.docx")
    args = ap.parse_args()

    data = json.loads(Path(args.data).read_text())
    report = fill_docx(args.docx, args.fields, data, args.output)
    print(f"Filled: {report['num_filled']}  Missing: {report['num_missing']}"
          f"  -> {args.output}")
